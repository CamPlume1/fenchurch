from io import BytesIO
import streamlit as st
import pandas as pd
from docx import Document
from perplexity import PerplexityResponse, query_preplexity

## Env
test_key = st.secrets["my_secrets"]["api_key"]

## Data
df = pd.read_excel("data_cleaned_3.xlsx", index_col=1)
df = df[~df.index.duplicated(keep='first')]
options = list(df.index.values)

## Asset manager prompts ##
def firm_description_general(firm_name: str) -> str:
    return f"Give me a 5-10 sentence firm overview of {firm_name}. Be concise, and use bullet points. Don't give a greeting, salutation, lead in, or introduction. BULLET POINTS ONLY"

def brief_history(firm_name: str) -> str:
    return f"Provide a brief, concise history of the {firm_name} in bullet point form. Don't give a greeting, salutation, or introduction. BULLET POINTS ONLY"

def primary_offerings(firm_name: str)-> str:
    return f"What are the primary offerings of {firm_name}? Please provide a concise response in bullet point form"

def firm_size(firm_name: str) -> str:
    return f"What is the size of {firm_name}? Please provide a concise response in bullet point form. Don't give a greeting, salutation, or introduction. BULLET POINTS ONLY"

def leadership(firm_name: str) -> str:
    return f"Give me a brief, concise description of key people associated with the RIA {firm_name}. Focus especially on firm leadership."


## AI Setup ##

def wrap(prompt: str) -> PerplexityResponse:
    return query_preplexity(prompt, test_key)


## Download Processing ##
 # Generate DOCX and provide download button
def generate_docx(data: dict[str, PerplexityResponse], sec: pd.DataFrame, user_input: str) -> BytesIO:
    doc = Document()
    doc.add_heading(f"RIA Report: {user_input}", level=1)

    doc.add_heading("SEC Disclosed Data - as of 12/02/2024", level=2)
    paragraph = doc.add_paragraph()

    # Home office
    state = f"{sec.at[user_input, 'Main Office State']}, " if pd.notna(sec.at[user_input, 'Main Office State']) else ""
    if pd.notna(sec.at[user_input, 'Main Office Country']):
        paragraph.add_run("Home Office: ").bold = True
        paragraph.add_run(f"{sec.at[user_input, 'Main Office City']}, {state}{sec.at[user_input, 'Main Office Country']}\n")
    else: paragraph.add_run("Home Office Not Disclosed, May Be Present in AI Research\n").bold = True

    paragraph.add_run("Discretionary AUM: ").bold = True
    paragraph.add_run(f"${sec.at[user_input, 'Discretionary AUM']}\n")

    paragraph.add_run("Non-Discretionary AUM: ").bold = True
    paragraph.add_run(f"${sec.at[user_input, 'Non-Discretionary AUM']}\n")

    paragraph.add_run("Total AUM: ").bold = True
    paragraph.add_run(f"${sec.at[user_input, 'Total AUM']}\n")

    paragraph.add_run("FTEs: ").bold = True
    paragraph.add_run(f"{sec.at[user_input, '#FTEs']}\n")

    paragraph.add_run("Number of Investment Prof: ").bold = True
    paragraph.add_run(f"{sec.at[user_input, '#Invest. Prof.']}\n")


    doc.add_heading("AI Generated Firm Research", level=2)
    for key, value in data.items():
        doc.add_heading(key.capitalize(), level=3)
        doc.add_paragraph(value.getText())
        
        # Create a new paragraph for citations
        paragraph = doc.add_paragraph()
        bold_run = paragraph.add_run("Citations: \n")  # Add "Citations" in bold
        bold_run.bold = True
        paragraph.add_run(value.getCitations())  # Append citations normally

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to convert a dictionary to Excel in-memory
def to_excel(a_dict) -> BytesIO:
    a_df = pd.DataFrame(a_dict, index=[0])
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        a_df.to_excel(writer, index=False, sheet_name='Sheet1')
    processed_data = output.getvalue()
    return processed_data

## Web App ##
st.title("RIA Reports")

# Password input
password = st.text_input("Enter password", type="password")

# User input for the firm name
user_input = st.selectbox("Enter firm name:", options)


# Initialize data accumulator
data_acc = {}

# Handle report generation logic
if st.button("Generate RIA report"):

    # Add section for SEC data
    st.header("Reported SEC Data - as of 12/02/2024")

    state = f"{df.at[user_input, 'Main Office State']}, " if pd.notna(df.at[user_input, 'Main Office State']) else ""

    if pd.notna(df.at[user_input, 'Main Office Country']):
        st.markdown(f"**Home office location:** {df.at[user_input, 'Main Office City']}, {state}{df.at[user_input, 'Main Office Country']}")
    else:
        st.markdown("**Home Office Location not reported, may be present in AI generated research**")
    st.markdown("#### AUM")
    st.markdown(f"**Discretionary AUM:** ${df.at[user_input, "Discretionary AUM"]:,.0f}")
    st.markdown(f"**Non-Discretionary AUM:** ${df.at[user_input, "Non-Discretionary AUM"]:,.0f}")
    st.markdown(f"**Total AUM:** ${df.at[user_input, "Total AUM"]:,.0f}")
    st.markdown("#### Staffing")
    st.markdown(f"**FTE's:** {df.at[user_input, "#FTEs"]}")
    st.markdown(f"**Number of Investment Prof.:** {df.at[user_input, "#Invest. Prof."]}")


    # Add a section heading
    st.header("Firm Research: AI Generated")

    # Validate password
    if password == st.secrets['my_secrets']['password']:
        if user_input:
            # Firm Overview
            data_acc['overview'] = wrap(firm_description_general(user_input))
            st.markdown("#### Firm Overview: AI GENERATED")
            st.text(f"{data_acc['overview'].getText()}")
            st.markdown(f"**Citations:** \n{data_acc['overview'].getMarkdownCitations()}", unsafe_allow_html=True)

            # Brief History
            data_acc['history'] = wrap(brief_history(user_input))
            st.markdown("#### Firm History: AI GENERATED")
            st.text(f"{data_acc['history'].getText()}")
            st.markdown(f"**Citations:** \n{data_acc['history'].getMarkdownCitations()}", unsafe_allow_html=True)

            # Primary Offerings
            data_acc['offerings'] = wrap(primary_offerings(user_input))
            st.markdown("#### Primary Offerings: AI GENERATED")
            st.text(f"{data_acc['offerings'].getText()}")
            st.markdown(f"**Citations:** \n{data_acc['offerings'].getMarkdownCitations()}", unsafe_allow_html=True)


            # Firm Size
            data_acc['size'] = wrap(firm_size(user_input))
            st.markdown("#### Firm Size: AI GENERATED")
            st.text(f"{data_acc['size'].getText()}")
            st.markdown(f"**Citations:** \n{data_acc['size'].getMarkdownCitations()}", unsafe_allow_html=True)

            # Firm Leadership
            data_acc['leadership'] = wrap(leadership(user_input))
            st.markdown("#### Firm Leadership: AI GENERATED")
            st.text(f"{data_acc['leadership'].getText()}")
            st.markdown(f"**Citations:** \n{data_acc['leadership'].getMarkdownCitations()}", unsafe_allow_html=True)

            docx_buffer = generate_docx(data_acc, df, user_input)

            st.download_button(
                label="Download Report as DOCX",
                data=docx_buffer,
                file_name=f"{user_input}_RIA_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            st.download_button(
                label="Download Report as XLSX: In Progress",
                data=to_excel(data_acc),
                file_name=f"{user_input}_RIA_Report.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
        else:
            st.error("Please enter a firm name to generate the report.")
    else:
        st.error("Incorrect password. Please try again.")
